import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Row
from docx.text.font import Font
from docx.text.paragraph import Paragraph
import copy

from token_docx import Tokens, TokenTypeString, TokenTypeCollection, SupportTable, get_font_and_size, \
    set_font_and_size


class DocxAnalyzer:
    __document: Document
    __tokens: Tokens
    __data: any

    def __init__(self):
        self.__document = None
        self.__tokens = Tokens()
        self.__data = None

    # Открытие шаблона документа
    def open_document(self, name_file: str):
        self.__document = Document(name_file)

    # Начало поиска токенов в таблицах
    def __search_token_tables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    sup_table = SupportTable()
                    sup_table.table = table
                    sup_table.row = row
                    sup_table.old_cells = copy.deepcopy(row.cells)
                    self.__search_token_paragraph(cell.paragraphs, sup_table)

    # Начало поиска токенов
    def start_search_tokens(self):
        tables = self.__document.tables
        if len(tables) != 0:
            self.__search_token_tables(tables)
        paragraphs = self.__document.paragraphs
        self.__search_token_paragraph(paragraphs)

    # Анализ и определение типов токенов
    def __analyze_token_type(self, paragraph: Paragraph, is_table=None):
        for word in paragraph.text.split():
            if "{{" and "}}" in word:
                token = re.search('{{(.+?)}}', word).group(0)
                if "." in token:
                    self.__tokens.add_token_type_collection(token, token, paragraph, is_table)
                else:
                    style = paragraph.style.name
                    if 'List' in style:
                        self.__tokens.add_token_type_list(token, paragraph)
                    else:
                        self.__tokens.add_token_type_string(token, paragraph)

    # Поиск токенов в параграфе
    def __search_token_paragraph(self, paragraphs, table=None):
        for paragraph in paragraphs:
            if "{{" and "}}" in paragraph.text:
                self.__analyze_token_type(paragraph, table)

    def init_data(self, data):
        self.__data = data

    # Анализ данных типа строка
    def __replace_str(self, key, value):
        temp_token: TokenTypeString = self.__tokens.TokensTypeString.get(key)
        if temp_token is None:
            return
        temp_token.paragraph = self.__replace_str_in_paragraph(temp_token.paragraph, temp_token.main_token, value,
                                                               temp_token.font)

    # Замена текста на новый в параграфе
    def __replace_str_in_paragraph(self, paragraph, token: str, new_name: str, font: Font):
        # clear_token = token.translate({ord(i): None for i in '{}'})
        # font = get_font_and_size(paragraph.runs, clear_token)
        paragraph.text = paragraph.text.replace(token, new_name)
        set_font_and_size(paragraph.runs, new_name, font)
        return paragraph

    # Восстановление коллекции
    def __restoring_collection(self, last_token, token_collection: TokenTypeCollection):
        copy_last_token = copy.copy(last_token)
        for token in token_collection.sub_tokens:
            value = token_collection.sub_tokens.get(token)
            if value is None:
                continue
            type_token = type(value)
            if type_token == TokenTypeString:
                if (value.old_text_paragraph != copy_last_token.old_text_paragraph
                        or (value.old_text_paragraph == last_token.old_text_paragraph and last_token != value)
                        or (
                                value.old_text_paragraph == last_token.old_text_paragraph and copy_last_token.main_token == value.main_token)):
                    copy_last_token.paragraph = self.__insert_paragraph_after(copy_last_token.paragraph,
                                                                              value.old_text_paragraph, value.paragraph,
                                                                              value.font, value.format)
                    copy_last_token.main_token = value.main_token
                    copy_last_token.old_text_paragraph = copy_last_token.paragraph.text
                value.paragraph = copy_last_token.paragraph
            else:
                self.__restoring_collection(copy_last_token, value)

    # Вставка новой строки в таблицу
    def __insert_after_row(self, table, ix):
        tbl = table._tbl
        successor = tbl.tr_lst[ix]
        tr = tbl._new_tr()
        for gridCol in tbl.tblGrid.gridCol_lst:
            tc = tr.add_tc()
            tc.width = gridCol.w
        successor.addprevious(tr)
        temp = table.rows[ix]
        return temp

    # Восстановление коллекции в таблицах
    def __update_collection_in_table(self, token_collection: TokenTypeCollection):
        for cell in token_collection.table.row.cells:
            self.__search_token_paragraph(cell.paragraphs)
        for token in token_collection.sub_tokens:
            value = token_collection.sub_tokens.get(token)
            if value is None:
                continue
            type_token = type(value)
            if type_token == TokenTypeString:
                continue
            else:
                self.__update_collection_in_table(value)

    # Выставление цвета заливки ячейки в таблице
    def __set_cell_bg_color(self, parent_cell, child_cell):
        pattern = re.compile('w:fill=\"(\S*)\"')
        match = pattern.search(parent_cell._tc.xml)
        if match is not None:
            result = match.group(1)
            tcPr = child_cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), result)
            tcPr.append(shd)

    # Создание новой строки в таблице
    def __create_new_row(self, cells, table: Table, height: int):
        new_row: _Row = table.add_row()
        new_row.height = height
        for i, index in enumerate(cells):
            for j, paragraph in enumerate(index.paragraphs):
                font = get_font_and_size(paragraph.runs, paragraph.text)
                new_row.cells[i].paragraphs[j].text = paragraph.text
                new_row.cells[i].paragraphs[j].style = paragraph.style

                new_row.cells[i].vertical_alignment = cells[i].vertical_alignment
                self.__set_cell_bg_color(cells[i], new_row.cells[i])

                set_font_and_size(new_row.cells[i].paragraphs[j].runs,
                                  paragraph.text, font)
                self.__set_paragraph_format(new_row.cells[i].paragraphs[j].
                                            paragraph_format, paragraph.paragraph_format)
        return new_row

    # Анализ данных типа dict для заполнения токенов
    def __replace_dict(self, main_key: str, data, collection=None):
        temp_token_collection = collection
        if collection is None:
            temp_token_collection: TokenTypeCollection = self.__tokens.TokensTypeCollection.get(main_key)

        if temp_token_collection is None:
            return

        last_token = None
        for key in data:
            value = data.get(key)
            type_value = type(value)
            if type_value == str:
                sub_token: TokenTypeString = temp_token_collection.sub_tokens.get(key)
                if sub_token is None:
                    continue
                self.__replace_str_in_paragraph(sub_token.paragraph, sub_token.main_token, value, sub_token.font)
                last_token = sub_token
            elif type_value == dict:
                sub_token: TokenTypeCollection = temp_token_collection.sub_tokens.get(key)
                if sub_token is None:
                    continue
                self.__replace_dict(main_key, value, sub_token)
            else:
                temp_type = type(value[0])
                if temp_type == dict:
                    sub_token: TokenTypeCollection = temp_token_collection.sub_tokens.get(key)
                    if sub_token is None:
                        continue
                    last_token = self.__replace_list_dict(key, value, sub_token)
                elif temp_type == str:
                    sub_token: TokenTypeString = temp_token_collection.sub_tokens.get(key)
                    if sub_token is None:
                        continue
                    last_token = self.__replace_list(main_key, value, sub_token)
        return last_token

    def __all_same(self, data):
        for key in data:
            value = data.get(key)
            now_type = type(value)
            if str != now_type:
                return False
        return True

    def __restoring_collection_in_table(self, token_collection: TokenTypeCollection):
        token_collection.table.row = self.__create_new_row(token_collection.table.old_cells,
                                                           token_collection.table.table,
                                                           token_collection.table.row.height)
        for token in token_collection.sub_tokens:
            value = token_collection.sub_tokens.get(token)
            if value is None:
                continue
            type_token = type(value)
            if type_token == TokenTypeString:
                continue
            else:
                self.__restoring_collection_in_table(value)

    # Анализ списка данных типа dict
    def __replace_list_dict(self, main_key: str, list_data: list, collection=None):
        i = 1
        temp_token_collection = collection
        if collection is None:
            temp_token_collection: TokenTypeCollection = self.__tokens.TokensTypeCollection.get(main_key)

        if temp_token_collection is None:
            return

        last_token = None
        simple_repeat = self.__all_same(list_data[0])
        for data in list_data:
            if i < len(list_data):
                be_more = True
                if simple_repeat and temp_token_collection.table:
                    temp_token_collection.table.row = self.__create_new_row(temp_token_collection.table.row.cells,
                                                                            temp_token_collection.table.table,
                                                                            temp_token_collection.table.row.height)
            else:
                be_more = False
            last_token = self.__replace_dict(main_key, data, temp_token_collection)
            if temp_token_collection.table and not simple_repeat and i < len(list_data):
                self.__restoring_collection_in_table(temp_token_collection)
            if be_more and temp_token_collection.table:
                self.__update_collection_in_table(temp_token_collection)
            elif be_more and last_token is not None and temp_token_collection.table is None:
                self.__restoring_collection(last_token, temp_token_collection)
            i = i + 1
        return last_token

    def __list_number(self, par, prev=None, level=None, num=True):
        xpath_options = {
            True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
            False: {'single': '', 'level': level},
        }

        def style_xpath(prefer_single=True):
            """
            The style comes from the outer-scope variable ``par.style.name``.
            """
            style = par.style.style_id
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
                ']/@w:abstractNumId'
            ).format(style=style, **xpath_options[prefer_single])

        def type_xpath(prefer_single=True):
            """
            The type is from the outer-scope variable ``num``.
            """
            type = 'decimal' if num else 'bullet'
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
                ']/@w:abstractNumId'
            ).format(type=type, **xpath_options[prefer_single])

        def get_abstract_id():
            """
            Select as follows:

                1. Match single-level by style (get min ID)
                2. Match exact style and level (get min ID)
                3. Match single-level decimal/bullet types (get min ID)
                4. Match decimal/bullet in requested level (get min ID)
                3. 0
            """
            for fn in (style_xpath, type_xpath):
                for prefer_single in (True, False):
                    xpath = fn(prefer_single)
                    ids = numbering.xpath(xpath)
                    if ids:
                        return min(int(x) for x in ids)
            return 0

        if (prev is None or
                prev._p.pPr is None or
                prev._p.pPr.numPr is None or
                prev._p.pPr.numPr.numId is None):
            if level is None:
                level = 0
            numbering = self.__document.part.numbering_part.numbering_definitions._numbering
            # Compute the abstract ID first by style, then by num
            anum = get_abstract_id()
            # Set the concrete numbering based on the abstract numbering ID
            num = numbering.add_num(anum)
            # Make sure to override the abstract continuation property
            num.add_lvlOverride(ilvl=level).add_startOverride(1)
            # Extract the newly-allocated concrete numbering ID
            num = num.numId
        else:
            if level is None:
                level = prev._p.pPr.numPr.ilvl.val
            # Get the previous concrete numbering ID
            num = prev._p.pPr.numPr.numId.val
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

    # Возвращает уровень(глубину) параграфа
    def __get_level_paragraph(self, paragraph: Paragraph):
        return paragraph._p.pPr.numPr.ilvl.val

    # Выставление формата параграфа
    def __set_paragraph_format(self, format: Paragraph.paragraph_format, new_format: Paragraph.paragraph_format):
        format.alignment = new_format.alignment
        format.first_line_indent = new_format.first_line_indent
        format.keep_together = new_format.keep_together
        format.keep_with_next = new_format.keep_with_next
        format.left_indent = new_format.left_indent
        format.line_spacing = new_format.line_spacing
        format.line_spacing_rule = new_format.line_spacing_rule
        format.page_break_before = new_format.page_break_before
        format.right_indent = new_format.right_indent
        format.space_after = new_format.space_after
        format.space_before = new_format.space_before
        if len(format.tab_stops) != 0:
            format.tab_stops = new_format.tab_stops
        format.widow_control = new_format.widow_control

    def __insert_paragraph_after(self, paragraph, text=None, repeat_paragraph=None, font=None, format=None):
        """Вставка нового параграфа после указанного"""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        level = None
        if text:
            new_para.add_run(text)
        if repeat_paragraph:
            new_para.style = repeat_paragraph.style
            set_font_and_size(new_para.runs, text, font)
            level = self.__get_level_paragraph(repeat_paragraph)
        if format:
            self.__set_paragraph_format(new_para.paragraph_format, format)
        self.__list_number(new_para, paragraph, level)
        return new_para

    def __replace_list(self, main_token, values, collection=None):
        if collection is None:
            temp_token: TokenTypeString = self.__tokens.TokensTypeString.get(main_token)
        else:
            temp_token: TokenTypeString = collection
        if temp_token is None:
            return
        i = 1
        for value in values:
            new_paragraph = None
            if i < len(values):
                new_paragraph = self.__insert_paragraph_after(temp_token.paragraph, temp_token.paragraph.text,
                                                              temp_token.paragraph, temp_token.font, temp_token.format)
            self.__replace_str_in_paragraph(temp_token.paragraph, temp_token.main_token, value, temp_token.font)
            if new_paragraph is not None:
                temp_token.paragraph = new_paragraph
            i = i + 1
        return temp_token

    def __analyze_type_in_list(self, main_key, list_data):
        temp_type = type(list_data[0])
        if temp_type == dict:
            self.__replace_list_dict(main_key, list_data)
        elif temp_type == str:
            self.__replace_list(main_key, list_data)

    # Запуск анализа данных и замены токенов
    def start_replace(self):
        for key in self.__data:
            value = self.__data.get(key)
            type_value = type(value)
            if type_value == str:
                self.__replace_str(key, value)
            elif type_value == dict:
                self.__replace_dict(key, value)
            elif type_value == list:
                self.__analyze_type_in_list(key, value)

    # Сохранение документа
    def save_file(self, new_name_file):
        self.__document.save(new_name_file)
