from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from HelpData import AuxiliaryToken
from HelpData import *


class DocxAnalyzer:
    TokensStyleDict = {}
    Data = {}
    ParagraphTokens = []
    TokenCollection = Tokens()

    def OpenDocument(self, nameDoc):
        """Открытие документа"""
        self.document = Document(nameDoc)

    def SearchTableToken(self, tables):
        """Поиск токенов в таблицах"""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self.SearchToken(cell.paragraphs)

    def WorkWithTokenTypeCollection(self, token):
        tokens = token.split('.')
        size = len(tokens)
        main_token = tokens[size - 1]
        sub_token = tokens[size - 2]

        pass

    def AnalyzeTypeToken(self, token, style):
        if '.' in token:
            self.WorkWithTokenTypeCollection(token)
        else:
            pass

        pass

    def add_token_string(self, paragraph: Paragraph):
        for word in paragraph.text.split():
            if "{{" and "}}" in word:
                self.TokenCollection.add_token_type_string(word, paragraph)



        # auxiliaryToken = AuxiliaryToken()
        # auxiliaryToken.SetParagraph(paragraph)
        # for word in paragraph.text.split():
        #     if "{{" and "}}" in word:
        #         style = paragraph.style.name
        #         temp_word = word.translate({ord(i): None for i in '{}'})
        #         auxiliaryToken.AddToken(temp_word, style)
        # self.ParagraphTokens.append(auxiliaryToken)

    def SearchToken(self, paragraphs, type=None):
        """Поиск токенов в параграфах"""
        for paragraph in paragraphs:
            for word in paragraph.text.split():
                if "{{" and "}}" in word:
                    # self.AddAuxiliaryToken(paragraph)
                    paragraphs.next()
                    # style = paragraph.style
                    # tempWord = word.translate({ord(i): None for i in '{}'})
                    # self.TokensStyleDict.update({tempWord: style})

    def AnalyzeData(self, data: dict):
        for auxiliaryToken in self.ParagraphTokens:
            for token in auxiliaryToken.tokens:
                temp_data = data.get(token)

    def list_number(self, par, prev=None, level=None, num=True):

        xpath_options = {
            True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
            False: {'single': '', 'level': level},
        }

        def style_xpath(prefer_single=True):
            """
            The style comes from the outer-scope variable ``par.style.name``.
            """
            style = par.style.style_id
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
                ']/@w:abstractNumId'
            ).format(style=style, **xpath_options[prefer_single])

        def type_xpath(prefer_single=True):
            """
            The type is from the outer-scope variable ``num``.
            """
            type = 'decimal' if num else 'bullet'
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
                ']/@w:abstractNumId'
            ).format(type=type, **xpath_options[prefer_single])

        def get_abstract_id():
            """
            Select as follows:

                1. Match single-level by style (get min ID)
                2. Match exact style and level (get min ID)
                3. Match single-level decimal/bullet types (get min ID)
                4. Match decimal/bullet in requested level (get min ID)
                3. 0
            """
            for fn in (style_xpath, type_xpath):
                for prefer_single in (True, False):
                    xpath = fn(prefer_single)
                    ids = numbering.xpath(xpath)
                    if ids:
                        return min(int(x) for x in ids)
            return 0

        if (prev is None or
                prev._p.pPr is None or
                prev._p.pPr.numPr is None or
                prev._p.pPr.numPr.numId is None):
            if level is None:
                level = 0
            numbering = self.document.part.numbering_part.numbering_definitions._numbering
            # Compute the abstract ID first by style, then by num
            anum = get_abstract_id()
            # Set the concrete numbering based on the abstract numbering ID
            num = numbering.add_num(anum)
            # Make sure to override the abstract continuation property
            num.add_lvlOverride(ilvl=level).add_startOverride(1)
            # Extract the newly-allocated concrete numbering ID
            num = num.numId
        else:
            if level is None:
                level = prev._p.pPr.numPr.ilvl.val
            # Get the previous concrete numbering ID
            num = prev._p.pPr.numPr.numId.val
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

    def insert_paragraph_after(self, paragraph, text=None):
        """Вставка нового параграфа после указанного"""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)

        self.list_number(new_para, paragraph)
        return new_para

    def StartSearch(self):
        """Начало поиска токенов"""
        tables = self.document.tables
        if (len(tables) != 0):
            self.SearchTableToken(tables)
        paragraphs = self.document.paragraphs
        self.SearchToken(paragraphs)

    def GetTokensStyleDict(self):
        """Возвращение словаря вида - {токен: стиль}"""
        return self.TokensStyleDict

    def init_data(self, data):
        """Инициализация словаря вида - {Токен: значение}"""
        self.Data = data

    def analyze_str(self, data: str):
        value = self.Data.get(data)
        TTS = self.TokenCollection.TokensTypeString.get(data)
        paragraph = TTS.paragraph
        font = self.GetFontAndSize(paragraph.runs, data)
        paragraph.text = paragraph.text.replace(data, value)
        self.SetFontAndSize(paragraph.runs, value, font)

    def analyze_dict(self, data: dict):
        pass

    def analyze_list(self, data: list):
        for sub_data in data:
            stype = type(sub_data)
            if stype == dict:
                pass
            elif stype == str:
                pass
            elif stype == list:
                self.analyze_list(sub_data)

    def start_replace(self):
        for sub_data in self.Data:
            stype = type(sub_data)
            if stype == list:
                self.analyze_list(sub_data)
            elif stype == str:
                self.analyze_str(sub_data)
            else:
                pass


    def new_analyze_list(self, paragraph, tokens):
        i = 0
        if len(tokens) > 1:
            parText = paragraph.text
        for token in tokens:
            if i == 0:
                pass

        pass

    #
    # под каждый тип данных свой массив
    # под список - коллекцию
    # под стоки - стринг
    #
    # каждый токен разбирать на свои массивы
    #
    # {{name.first}} {{name.second}} == {{firstName}} {{secondName}}

    def StartReplace(self):
        """Начало операции по замене токенов"""
        for auxiliaryToken in self.ParagraphTokens:
            if "list" in auxiliaryToken.paragraph.style.name:
                self.new_analyze_list(auxiliaryToken.paragraph, auxiliaryToken.tokens)
            else:
                pass

        # tables = self.document.tables
        # if (len(tables) != 0):
        #     self.ReplaceTable(tables)
        # paragraphs = self.document.paragraphs
        # self.ReplaceToken(paragraphs)

    def ReplaceTable(self, tables):
        """Замена токенов в таблицах"""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self.ReplaceToken(cell.paragraphs)

    def GetFontAndSize(self, runs, word):
        for run in runs:
            if word in run.text:
                return (run.font)

    def SetFontAndSize(self, runs, word, font):
        for run in runs:
            if word in run.text:
                run.font.name = font.name
                run.font.size = font.size

    def AnalyzeList(self, paragraph, word, value):
        """ """
        i = 0
        for tempWord in value:
            if i == 0:
                font = self.GetFontAndSize(paragraph.runs, word.translate({ord(i): None for i in '{}'}))
                paragraph.text = paragraph.text.replace(word, value[0])
                self.SetFontAndSize(paragraph.runs, tempWord, font)
            else:
                self.insert_paragraph_after(paragraph, tempWord)
            i = i + 1

    def ReplaceToken(self, paragraphs):
        """Замена токенов"""

        # for paragraph in paragraphs:
        #     for word in paragraph.text.split():
        #         if "{{" and "}}" in word:
        #             tempWord = word.translate({ord(i): None for i in '{}'})
        #             tempStyle = self.TokensStyleDict.get(tempWord)
        #             tempValue = self.TokenDataDict.get(tempWord)
        #             if 'List' in tempStyle.name:
        #                 self.AnalyzeList(paragraph, word, tempValue)

    def SaveDocument(self, newName):
        """Сохранение документа"""
        self.document.save(newName)
