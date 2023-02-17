from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import copy

from HelpData import Tokens, TokenTypeString, TokenTypeCollection, SimpleParagraphData, SupportTable


def get_font_and_size(runs, word):
    for run in runs:
        if word in run.text:
            return run.font


def set_font_and_size(runs, word, font):
    for run in runs:
        if word in run.text:
            run.font.name = font.name
            run.font.size = font.size


class newDocxAnalyzer:
    document: Document
    tokens: Tokens
    data: any
    paragraphs_in_work = []

    def open_document(self, name_file: str):
        self.document = Document(name_file)
        self.tokens = Tokens()

    def search_token_tables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    sup_table = SupportTable()
                    sup_table.table = table
                    sup_table.last_row = row
                    sup_table.old_row = copy.copy(row)
                    self.search_token_paragraph(cell.paragraphs, sup_table)

    def start_search_tokens(self):
        tables = self.document.tables
        if len(tables) != 0:
            self.search_token_tables(tables)
        paragraphs = self.document.paragraphs
        self.search_token_paragraph(paragraphs)

    def analyze_token_type(self, paragraph: Paragraph, is_table=None):
        for word in paragraph.text.split():
            if "{{" and "}}" in word:
                if "." in word:
                    self.tokens.add_token_type_collection(word, word, paragraph, is_table)
                else:
                    style = paragraph.style.name
                    if 'List' in style:
                        self.tokens.add_token_type_list(word, paragraph)
                    else:
                        self.tokens.add_token_type_string(word, paragraph)

    def search_token_paragraph(self, paragraphs, table=None):
        for paragraph in paragraphs:
            if "{{" and "}}" in paragraph.text:
                self.analyze_token_type(paragraph, table)

    def init_data(self, data):
        self.data = data

    def replace_str(self, key, value):
        temp_token: TokenTypeString = self.tokens.TokensTypeString.get(key)
        temp_token.paragraph = self.replace_str_in_paragraph(temp_token.paragraph, temp_token.main_token, value)

        # font = get_font_and_size(temp_token.paragraph.runs, key)
        # temp_token.paragraph.text = temp_token.paragraph.text.replace(temp_token.main_token, value)
        # set_font_and_size(temp_token.paragraph.runs, value, font)

    def replace_str_in_paragraph(self, paragraph, token, new_name):
        clear_token = token.translate({ord(i): None for i in '{}'})
        font = get_font_and_size(paragraph.runs, clear_token)
        paragraph.text = paragraph.text.replace(token, new_name)
        set_font_and_size(paragraph.runs, new_name, font)
        return paragraph

    def restoring_collection(self, last_token, token_collection: TokenTypeCollection):
        copy_last_token = copy.copy(last_token)
        for token in token_collection.sub_tokens:
            value = token_collection.sub_tokens.get(token)
            type_token = type(value)
            if type_token == TokenTypeString:
                if value.old_text_paragraph != copy_last_token.old_text_paragraph or last_token == value:
                    copy_last_token.paragraph = self.insert_paragraph_after(copy_last_token.paragraph,
                                                                            value.old_text_paragraph, value.paragraph)
                    copy_last_token.old_text_paragraph = copy_last_token.paragraph.text
                value.paragraph = copy_last_token.paragraph
            else:
                self.restoring_collection(copy_last_token, value)

    def insert_after_row(self, table, ix):
        tbl = table._tbl
        successor = tbl.tr_lst[ix]
        tr = tbl._new_tr()
        for gridCol in tbl.tblGrid.gridCol_lst:
            tc = tr.add_tc()
            tc.width = gridCol.w
        successor.addprevious(tr)
        return table.rows[ix]

    def restoring_collection_in_table(self, token_collection: TokenTypeCollection):
        for cell in token_collection.table.last_row.cells:
            self.search_token_paragraph(cell.paragraphs)

    def replace_list_dict(self, main_key, list_data, collection=None):
        i = 1
        temp_token_collection = collection
        if collection is None:
            temp_token_collection: TokenTypeCollection = self.tokens.TokensTypeCollection.get(main_key)
        last_token = None
        for data in list_data:
            if i < len(list_data):
                be_more = True
                if temp_token_collection.table:
                    index = temp_token_collection.table.last_row._index
                    temp_token_collection.table.last_row = self.insert_after_row(temp_token_collection.table.table,
                                                                                 index)
                    for i, index in enumerate(temp_token_collection.table.old_row.cells):
                        temp_token_collection.table.last_row.cells[i].paragraphs = copy.copy(index.paragraphs)
                        # for j, paragraph in enumerate(index.paragraphs):
                        #     temp_token_collection.table.last_row.cells[i] = copy.copy()
                        #     temp_token_collection.table.last_row.cells[i].paragraphs[0].style = paragraph.style

            else:
                be_more = False
            for key in data:
                value = data.get(key)
                type_value = type(value)
                if type_value == str:
                    sub_token: TokenTypeString = temp_token_collection.sub_tokens.get(key)
                    self.replace_str_in_paragraph(sub_token.paragraph, sub_token.main_token, value)
                    last_token = sub_token
                else:
                    temp_type = type(value[0])
                    if temp_type == dict:
                        sub_token: TokenTypeCollection = temp_token_collection.sub_tokens.get(key)
                        last_token = self.replace_list_dict(key, value, sub_token)
                    elif temp_type == str:
                        sub_token: TokenTypeString = temp_token_collection.sub_tokens.get(key)
                        last_token = self.replace_list(main_key, value, sub_token)
            if be_more and collection is None and temp_token_collection.table:
                self.restoring_collection_in_table(temp_token_collection)
            elif be_more and last_token is not None:
                self.restoring_collection(last_token, temp_token_collection)
            i = i + 1
        return last_token

    def list_number(self, par, prev=None, level=None, num=True):
        xpath_options = {
            True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
            False: {'single': '', 'level': level},
        }

        def style_xpath(prefer_single=True):
            """
            The style comes from the outer-scope variable ``par.style.name``.
            """
            style = par.style.style_id
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
                ']/@w:abstractNumId'
            ).format(style=style, **xpath_options[prefer_single])

        def type_xpath(prefer_single=True):
            """
            The type is from the outer-scope variable ``num``.
            """
            type = 'decimal' if num else 'bullet'
            return (
                'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
                ']/@w:abstractNumId'
            ).format(type=type, **xpath_options[prefer_single])

        def get_abstract_id():
            """
            Select as follows:

                1. Match single-level by style (get min ID)
                2. Match exact style and level (get min ID)
                3. Match single-level decimal/bullet types (get min ID)
                4. Match decimal/bullet in requested level (get min ID)
                3. 0
            """
            for fn in (style_xpath, type_xpath):
                for prefer_single in (True, False):
                    xpath = fn(prefer_single)
                    ids = numbering.xpath(xpath)
                    if ids:
                        return min(int(x) for x in ids)
            return 0

        if (prev is None or
                prev._p.pPr is None or
                prev._p.pPr.numPr is None or
                prev._p.pPr.numPr.numId is None):
            if level is None:
                level = 0
            numbering = self.document.part.numbering_part.numbering_definitions._numbering
            # Compute the abstract ID first by style, then by num
            anum = get_abstract_id()
            # Set the concrete numbering based on the abstract numbering ID
            num = numbering.add_num(anum)
            # Make sure to override the abstract continuation property
            num.add_lvlOverride(ilvl=level).add_startOverride(1)
            # Extract the newly-allocated concrete numbering ID
            num = num.numId
        else:
            if level is None:
                level = prev._p.pPr.numPr.ilvl.val
            # Get the previous concrete numbering ID
            num = prev._p.pPr.numPr.numId.val
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

    def get_level_paragraph(self, paragraph: Paragraph):
        return paragraph._p.pPr.numPr.ilvl.val

    def insert_paragraph_after(self, paragraph, text=None, repeat_paragraph=None):
        """Вставка нового параграфа после указанного"""
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        level = None
        if text:
            new_para.add_run(text)
        if repeat_paragraph:
            new_para.style = repeat_paragraph.style
            level = self.get_level_paragraph(repeat_paragraph)
        self.list_number(new_para, paragraph, level)
        return new_para

    def replace_list(self, main_token, values, collection=None):
        if collection is None:
            temp_token: TokenTypeString = self.tokens.TokensTypeString.get(main_token)
        else:
            temp_token: TokenTypeString = collection
        i = 1
        for value in values:
            new_paragraph = None
            if i < len(values):
                new_paragraph = self.insert_paragraph_after(temp_token.paragraph, temp_token.paragraph.text,
                                                            temp_token.paragraph)
            self.replace_str_in_paragraph(temp_token.paragraph, temp_token.main_token, value)
            if new_paragraph is not None:
                temp_token.paragraph = new_paragraph
            i = i + 1
        return temp_token

    def analyze_type_in_list(self, main_key, list_data):
        temp_type = type(list_data[0])
        if temp_type == dict:
            self.replace_list_dict(main_key, list_data)
        elif temp_type == str:
            self.replace_list(main_key, list_data)

    def start_replace(self):
        for key in self.data:
            value = self.data.get(key)
            type_value = type(value)
            if type_value == str:
                self.replace_str(key, value)
            elif type_value == list:
                self.analyze_type_in_list(key, value)

    def save_file(self, new_name_file):
        self.document.save(new_name_file)
